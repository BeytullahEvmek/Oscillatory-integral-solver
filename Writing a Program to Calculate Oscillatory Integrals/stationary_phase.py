# Libraries
import tkinter as tk #(For UI)
from tkinter import messagebox, filedialog # (For popups and to show where to save)
from sympy import symbols, diff, solve, sqrt, pi, exp, I, S, lambdify, E #(For math)
from sympy.parsing.sympy_parser import parse_expr
import numpy as np #(For numerical arrays)
from matplotlib.figure import Figure #(For Graph)
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from docx import Document #(For exporting to Docx)
from docx.shared import Inches
import os #(For saving the Docx correctly)

x, lam = symbols('x λ', real=True) #Idenifies the symbols

def stationary_phase_approx(f_expr, phi_expr, lam_val):
    phi_d = diff(phi_expr, x) # First derivative "φ'(x)"
    stat_points = solve(phi_d, x) # Solves φ'(x) = 0
    if not stat_points: # Aborts if there's nothing
        return None, None, "No stationary points found."

    if lam_val <= 0: #makes no input 0 and below not work
        messagebox.showerror("Invalid Input", "λ must be positive for the stationary phase approximation.")
        return None, None, "λ must be positive for the stationary phase approximation."

    x0 = stat_points[0] # Chooses the first stationary point
    phi_dd = diff(phi_expr, x, 2).subs(x, x0) # Second derivative "φ''(x₀)"
    f_val = f_expr.subs(x, x0)  # Evaluates f(x₀)
    phase_shift = S(pi)/4 if phi_dd > 0 else -S(pi)/4 # Choose phase shift ±π/4, if bigger than 0 it's +, if not it's-

    approx = f_val * sqrt(2 * pi / (lam * abs(phi_dd))) * exp(I * lam * phi_expr.subs(x, x0) + I * phase_shift) # Final approximation of I(λ)
    approx_real = approx.as_real_imag()[0]  # Takes the real part of I(λ)
    approx_at_lam = approx.subs(lam, lam_val).evalf() #putting the given λ
    approx_r = approx_at_lam.as_real_imag()[0]  # Extracts the real part of I(λ)
    approx_i = approx_at_lam.as_real_imag()[1]  # Extracts the imaginary part of I(λ)
    num_lam = lambdify(lam, approx_real, modules="numpy")  # Converts to a NumPy-friendly function of λ

    def pretty_expr(expr): # Makes so that the sympy terms turns into symbols
        s = str(expr)
        s = s.replace("pi", "π")
        s = s.replace("lambda", "λ")
        s = s.replace("**", "^")
        s = s.replace("sqrt", "√")
        s = s.replace('exp', 'e^')
        s = s.replace("I", "i")
        return s

    # Step-by-step solution explanation
    steps = [
        "Given inputs:",
        f"  f(x) = {pretty_expr(f_expr)}",
        f"  φ(x) = {pretty_expr(phi_expr)}",
        f"  λ = {pretty_expr(lam_val)}",
        "",
        "Step 1: Compute the derivative of φ(x):",
        "  General formula: dφ/dx = φ'(x)",
        f"  Computed: φ'(x) = {pretty_expr(phi_d)}",
        "",
        "Step 2: Find stationary points by solving φ'(x) = 0:",
        "  Solve: φ'(x) = 0",
        f"  Stationary points: {[pretty_expr(pt) for pt in stat_points]}",
        f"  Chosen stationary point: x₀ = {pretty_expr(x0)}",
        "",
        "Step 3: Compute the second derivative of φ(x) at x₀:",
        "  General formula: φ''(x₀) = d²φ/dx² evaluated at x₀",
        f"  Computed: φ''(x₀) = {pretty_expr(phi_dd)}",
        "",
        "Step 4: Evaluate f(x) at x₀:",
        "  General formula: f(x₀)",
        f"  Computed: f(x₀) = {pretty_expr(f_val)}",
        "",
        "Step 5: Write the stationary phase approximation formula:",
        "  General formula:",
        "    I(λ) ≈ f(x₀) * √(2π / (λ * |φ''(x₀)|)) * e^(iλφ(x₀) ± iπ/4)",
        "    where the sign ± depends on the sign of φ''(x₀):",
        "    + if φ''(x₀) > 0, use +π/4",
        "    - if φ''(x₀) < 0, use -π/4",
        "  Substituted expression:",
        f"    I(λ) ≈ {pretty_expr(f_val)} * √(2π / (λ * |{pretty_expr(phi_dd)}|)) * e^(iλ * {pretty_expr(phi_expr.subs(x, x0))} {'+ π/4' if phi_dd > 0 else '- π/4'})",
        "",
        "Step 6: Final computed approximation:",
        f"  I(λ) ≈ {'∞' if approx == S.ComplexInfinity else pretty_expr(approx)}"
        "",
        f"  I(λ) ≈ {pretty_expr(approx_at_lam)}",
        f"  Real part: {approx_r}",
        f"  Imaginary part: {approx_i}",
    ]

    return num_lam, steps, None # Returns the function of λ, steps, and None if there's no error

def export_to_docx(steps, fig, doc_path): #Exports the Docx file
    doc = Document()
    doc.add_heading('Stationary Phase Approximation', 0) #Makes a header inside Docx

    for step in steps: #Adds steps
        doc.add_paragraph(step)

    # Temporarily saves the graph to insert it into the Docx
    temp_path = 'temp_plot.png'
    fig.savefig(temp_path)
    doc.add_picture(temp_path, width=Inches(6)) # Preset width

    doc.save(doc_path) # Saves Docx document


    if os.path.exists(temp_path): # Cleans temporary files
        os.remove(temp_path)

    return doc_path

def show_stationary_phase(parent_frame, go_back): # GUI for the Stationary Phase page
    for widget in parent_frame.winfo_children(): # Clears previous widgets
        widget.destroy()

    parent_frame.configure(bg="#ccffcc")  # Sets background color
    # Title
    title = tk.Label(parent_frame, text="Stationary Phase Approximation", font=("Helvetica", 40, "bold"), bg="#ccffcc")
    title.grid(row=0, column=0, columnspan=4, pady=20)
    # Function input
    tk.Label(parent_frame, text="Enter f(x):", bg="#ccffcc", font=("Arial", 30)).grid(row=1, column=0, sticky='e', padx=10, pady=5)
    f_entry = tk.Entry(parent_frame, width=50, font=("Arial", 30))
    f_entry.grid(row=1, column=1, columnspan=3, padx=10, pady=5)
    f_entry.insert(0, "e^(-x^2)")
    # Phase input
    tk.Label(parent_frame, text="Enter φ(x):", bg="#ccffcc", font=("Arial", 30)).grid(row=2, column=0, sticky='e', padx=10, pady=5)
    phi_entry = tk.Entry(parent_frame, width=50, font=("Arial", 30))
    phi_entry.grid(row=2, column=1, columnspan=3, padx=10, pady=5)
    phi_entry.insert(0, "x^2")
    # Lambda input
    tk.Label(parent_frame, text="Enter λ:", bg="#ccffcc", font=("Arial", 30)).grid(row=3, column=0, sticky='e', padx=10, pady=5)
    lam_entry = tk.Entry(parent_frame, width=20, font=("Arial", 30))
    lam_entry.grid(row=3, column=1, columnspan=3, padx=10, pady=5)
    lam_entry.insert(0, "10")
    # Text output for the solution
    output_text = tk.Text(parent_frame, width=60, height=20, font=("Arial", 15))
    output_text.grid(row=4, column=0, columnspan=2, padx=15, pady=10, sticky='nsew')
    # Graph output
    fig = Figure(figsize=(6, 4), dpi=100)
    ax = fig.add_subplot(111)
    canvas = FigureCanvasTkAgg(fig, master=parent_frame)
    canvas.get_tk_widget().grid(row=4, column=2, columnspan=2, padx=15, pady=10, sticky='nsew')

    #Variables
    error_result = None
    steps_result = None

    def on_solve(): #What happens if the solve button is pressed
        nonlocal error_result, steps_result
        error_result = None
        steps_result = None

        output_text.delete("1.0", tk.END)  # Clear previous answer
        ax.clear()  # Clear previous graph

        # Gets
        f_input = f_entry.get().replace('^', '**')
        phi_input = phi_entry.get().replace('^', '**')


        # Parses input expressions with error handling
        try:
            f_expr = parse_expr(f_input, local_dict={'x': x, 'e': E})
            phi_expr = parse_expr(phi_input, local_dict={'x': x, 'e': E})
            lam_val = float(lam_entry.get())
        except Exception as e:
            messagebox.showerror("Input error", f"Error parsing input:\n{e}")
            return

        # Computes stationary phase approximation
        f_lam, steps, error = stationary_phase_approx(f_expr, phi_expr, lam_val)
        #Error failsafe
        if error:
            output_text.insert(tk.END, error)
            canvas.draw()
            error_result = error
            return

        # Display step-by-step solution
        steps_result = steps
        output_text.insert(tk.END, "\n".join(steps))

        # Checks if approximation is infinite, if so, no graph is generated
        approx_str = steps[-1]
        if "∞" in approx_str:
            ax.set_title("Approximation is infinite; no graph available")
            ax.set_xticks([])  # Remove x ticks
            ax.set_yticks([])  # Remove y ticks
            canvas.draw()
            return

        # Generate λ values for plotting
        lam_vals = np.linspace(lam_val - 15, lam_val + 15, 400)

        # Handles the errors may be caused by λ range
        try:
            y_vals = f_lam(lam_vals)
        except Exception as e:
            messagebox.showerror("Evaluation error", f"Error evaluating approximation:\n{e}")
            return

        #Evaluates λ safely
        try:
            y_at_lam = f_lam(lam_val)
        except Exception as e:
            messagebox.showerror("Input error", f"Error parsing λ:\n{e}")
            return

        # Graph
        ax.plot(lam_vals, y_vals, label="Approximation") #Plots the real part of the approximation over λ
        ax.axvline(lam_val, color='red', linestyle='--', label=f"λ = {lam_val}") # Highlights the λ
        ax.scatter([lam_val], [y_at_lam], color='red') # Marks the λ
        ax.set_title("Stationary Phase Approximation (Real Part)") #Title
        ax.set_xlabel("λ") #Whats on x axis
        ax.set_ylabel("Re(I(λ))") #Whats on y axis
        ax.legend()  #Adds a box that shows the details
        ax.grid(True) # Grids the graph
        canvas.draw() # Refreshes the graph

    def on_export():  #What happens if the export button is pressed
        # If there's an error, don't allow and show the corresponding popup
        if error_result is not None:
            messagebox.showerror("Error", error_result)
            return
        # No solution no export
        if not steps_result:
            messagebox.showerror("Export Error", "Please solve the problem first.")
            return

        # Allows to select a location and filename to save the Docx
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")],
            title="Save solution as..."
        )
        # If user cancels the export, aborts it
        if not file_path:
            return

        path = export_to_docx(steps_result, fig, file_path) #Exports everything into the path
        messagebox.showinfo("Export Success", f"Solution exported to:\n{path}") #Export completion pop up

    #Buttons
    solve_btn = tk.Button(parent_frame, text="Solve", font=("Arial", 18), command=on_solve)
    solve_btn.grid(row=5, column=0, pady=10, columnspan=2, sticky="ew", padx=5)

    export_btn = tk.Button(parent_frame, text="Save to Word", font=("Arial", 18), command=on_export)
    export_btn.grid(row=5, column=2, pady=10, sticky="ew", padx=5)

    back_btn = tk.Button(parent_frame, text="Back to Main", font=("Arial", 18),command=lambda: go_back(parent_frame))
    back_btn.grid(row=5, column=3, pady=10, sticky="ew", padx=5)

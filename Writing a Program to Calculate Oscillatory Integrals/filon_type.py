# Libraries
import tkinter as tk #(For UI)
from tkinter import messagebox, filedialog # (For popups and to show where to save)
from sympy import symbols, E # (For math)
from sympy.parsing.sympy_parser import parse_expr # Parses inputs safely
import numpy as np #(For numerical arrays)
from matplotlib.figure import Figure #(For Graph)
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from docx import Document  #(For exporting to Docx)
from docx.shared import Inches
import os #(For saving the Docx correctly)

x, lam = symbols('x λ', real=True) #Idenifies the symbols

def filon_approx(f_expr, a_val, b_val, lam_val):
    # Defines interpolation points
    x0 = a_val
    x1 = (a_val + b_val) / 2
    x2 = b_val
    #Evaluates f(x) at interpolation points
    f0 = f_expr.subs(x, x0)
    f1 = f_expr.subs(x, x1)
    f2 = f_expr.subs(x, x2)
    #Step size for Filon formula
    h = (x2 - x0) / 2
    #Defining Filon weight functions
    def A(z): #
        return (2*np.sin(z) - z*(np.cos(z) + 1)) / z**3 if z != 0 else 1/3

    def B(z):
        return (4*np.sin(z) - z*(np.cos(z) + 3)) / z**3 if z != 0 else 4/3

    def C(z):
        return (2*np.sin(z) - z*(3*np.cos(z) + 1)) / z**3 if z != 0 else 1/3

    #Calculating z (used on weights) (counts as the scaled frequency)
    z = lam_val * h
    # Converting symbolic function values to float
    f0_num = float(f0.evalf())
    f1_num = float(f1.evalf())
    f2_num = float(f2.evalf())
    # Computing the Filon approximation for the given λ
    approx = h * (A(z)*f0_num + B(z)*f1_num + C(z)*f2_num) * np.exp(1j * lam_val * x0) # Filon Formula

    # Vectorized approximation over a range of λ values
    def filon_vec(lambda_array):
        results = []
        for lam_local in lambda_array:
            z = lam_local * h     #Calculating z (used on weights)
            result = h * (A(z)*f0_num + B(z)*f1_num + C(z)*f2_num) * np.exp(1j * lam_local * x0) # Filon Formula
            results.append(result) #Appends result
        return np.array(results)
    # Pretty formatting for expressions in the output
    def pretty_expr(expr):
        s = str(expr)
        s = s.replace("pi", "π")
        s = s.replace("lambda", "λ")
        s = s.replace("**", "^")
        s = s.replace("sqrt", "√")
        s = s.replace('exp', 'e^')
        return s
    # Step-by-step solution explanation
    steps = [
        "Given inputs:",
        f"  f(x) = {pretty_expr(f_expr)}",
        f"  φ(x) = x (fixed for Filon method)",
        f"  Integration interval: [{a_val}, {b_val}]",
        "",
        "Step 1: Select three points for quadratic interpolation:",
        f"  x₀ = {x0}, x₁ = {x1}, x₂ = {x2}",
        "",
        "Step 2: Evaluate f(x) at these points:",
        f"  f(x₀) = {pretty_expr(f0)}",
        f"  f(x₁) = {pretty_expr(f1)}",
        f"  f(x₂) = {pretty_expr(f2)}",
        "",
        "Step 3: Compute weights A(z), B(z), C(z) with z = λ h:",
        "  A(z) = (2 sin z - z (cos z + 1)) / z³",
        "  B(z) = (4 sin z - z (cos z + 3)) / z³",
        "  C(z) = (2 sin z - z (3 cos z + 1)) / z³",
        "",
        "Step 4: Apply Filon approximation formula:",
        "  I(λ) ≈ h * (A(z) f(x₀) + B(z) f(x₁) + C(z) f(x₂)) * e^{i λ x₀}",
        f"  where z = {lam_val} * {h} = {z}",
        "",
        f"Step 5: Final approximation at λ = {lam_val}:",
        f"  I(λ) ≈ {str(approx).replace('j', 'i')}",
        f"  Real part: {approx.real}",
        f"  Imaginary part: {approx.imag}"
    ]

    return filon_vec, steps, None



def export_to_docx(steps, fig, doc_path): #Exports the Docx file
    doc = Document()
    doc.add_heading('Filon Quadrature Approximation', 0) #Makes a header inside Docx

    for step in steps:
        doc.add_paragraph(step)

    # Temporarily saves the graph to insert it into the Docx
    temp_path = 'temp_plot.png'
    fig.savefig(temp_path)
    doc.add_picture(temp_path, width=Inches(6))

    doc.save(doc_path)   # Saves Docx document

    if os.path.exists(temp_path): # Cleans temporary files
        os.remove(temp_path)

    return doc_path


def show_filon_type(parent_frame, go_back): # GUI logic for the Filon Function page
    for widget in parent_frame.winfo_children():# Clears previous widgets
        widget.destroy()

    parent_frame.configure(bg="#ccffcc")# Sets background color
    # Title
    title = tk.Label(parent_frame, text="Filon Quadrature Approximation", font=("Helvetica", 40, "bold"), bg="#ccffcc")
    title.grid(row=0, column=0, columnspan=4, pady=20)
    # Function input
    tk.Label(parent_frame, text="Enter f(x):", bg="#ccffcc", font=("Arial", 30)).grid(row=1, column=0, sticky='e', padx=10, pady=5)
    f_entry = tk.Entry(parent_frame, width=50, font=("Arial", 30))
    f_entry.grid(row=1, column=1, columnspan=3, padx=10, pady=5)
    f_entry.insert(0, "e^(-x^2)")
    # Phase label (it's preset nevertheless)
    tk.Label(parent_frame, text="φ(x) is fixed as x", bg="#ccffcc", font=("Arial", 28, "italic")).grid(row=2, column=0, columnspan=2, pady=10)
    #Lower limit input
    tk.Label(parent_frame, text="Enter integration lower limit a:", bg="#ccffcc", font=("Arial", 20)).grid(row=3, column=0, sticky='e', padx=10, pady=5)
    a_entry = tk.Entry(parent_frame, width=20, font=("Arial", 20))
    a_entry.grid(row=3, column=1, padx=10, pady=5)
    a_entry.insert(0, "0")
    #Upper limit input
    tk.Label(parent_frame, text="Enter integration upper limit b:", bg="#ccffcc", font=("Arial", 20)).grid(row=3, column=2, sticky='e', padx=10, pady=5)
    b_entry = tk.Entry(parent_frame, width=20, font=("Arial", 20))
    b_entry.grid(row=3, column=3, padx=10, pady=5)
    b_entry.insert(0, "1")
    # Lambda input
    tk.Label(parent_frame, text="Enter λ for approximation:", bg="#ccffcc", font=("Arial", 20)).grid(row=2, column=2, sticky='e', padx=10, pady=5)
    lam_entry = tk.Entry(parent_frame, width=15, font=("Arial", 20))
    lam_entry.grid(row=2, column=3, padx=10, pady=5)
    lam_entry.insert(0, "10")
    # Text output for the solution
    output_text = tk.Text(parent_frame, width=60, height=20, font=("Arial", 15))
    output_text.grid(row=4, column=0, columnspan=2, padx=15, pady=10, sticky='nsew')
    # Graph output
    fig = Figure(figsize=(6, 4), dpi=100)
    ax = fig.add_subplot(111)
    canvas = FigureCanvasTkAgg(fig, master=parent_frame)
    canvas.get_tk_widget().grid(row=4, column=2, columnspan=2, padx=15, pady=10, sticky='nsew')

    # Variables
    error_result = None
    steps_result = None
    def on_solve():  #What happens if the solve button is pressed
        nonlocal error_result, steps_result
        error_result = None
        steps_result = None


        output_text.delete("1.0", tk.END) # Clear previous answer
        ax.clear() # Clear previous graph

        # Get
        f_input = f_entry.get().replace('^', '**')
        # Parses input expressions with error handling
        try:
            f_expr = parse_expr(f_input, local_dict={'x': x, 'e': E})
        except Exception as e:
            error_result = f"Error parsing input:\n{e}"
            messagebox.showerror("Input error", error_result)
            return
        # Parses a & b and checks if a is smaller than b
        try:
            a_val = float(a_entry.get())
            b_val = float(b_entry.get())
            if b_val <= a_val:
                raise ValueError("Upper limit must be greater than lower limit")
        except Exception as e:
            error_result = f"Error in integration limits:\n{e}"
            messagebox.showerror("Input error", error_result)
            return
        # Parses λ input safely
        try:
            lam_val = float(lam_entry.get())
            if lam_val == 0:
                raise ValueError("Program not capable of handling λ = 0.")
        except Exception as e:
            error_result = f"Error parsing λ:\n{e}"
            messagebox.showerror("Input error", error_result)
            return

        # Computes filon approximation
        filon_func, steps, error = filon_approx(f_expr, a_val, b_val, lam_val)
        # Error failsafe
        if error:
            error_result = error
            output_text.insert(tk.END, error_result)
            canvas.draw()
            return

        # Display step-by-step solution
        steps_result = steps
        output_text.insert(tk.END, "\n".join(steps_result))

        # Plot over a range of λ values
        lam_vals = np.linspace(-20, 20, 400)
        try:
            y_vals = np.real(filon_func(lam_vals))
        except Exception as e:
            error_result = f"Error evaluating approximation:\n{e}"
            messagebox.showerror("Evaluation error", error_result)
            return

        #Evaluates λ safely
        try:
            y_at_lam = np.real(filon_func([lam_val])[0])
        except Exception as e:
            messagebox.showerror("Input error", f"Error evaluating approximation at λ:\n{e}")
            return

        ax.plot(lam_vals, y_vals, label="Re(I(λ))") #Plots the real part of the approximation over λ
        ax.axvline(lam_val, color='red', linestyle='--', label=f"λ = {lam_val}")# Highlights the λ
        ax.scatter([lam_val], [y_at_lam], color='red') # Marks the λ
        ax.set_title("Filon Quadrature Approximation (Real part)") #Title
        ax.set_xlabel("λ") #Whats on x axis
        ax.set_ylabel("Re(I(λ))") #Whats on y axis
        ax.grid(True)  # Grids the graph
        ax.legend()  #Adds a box that shows the details
        canvas.draw() # Refreshes the graph


    def on_export(): #What happens if the export button is pressed
        nonlocal error_result, steps_result
        # If there's an error, don't allow and show the corresponding popup
        if error_result is not None:
            messagebox.showerror("Error", error_result)
            return
        # No solution no export
        if not steps_result:
            messagebox.showerror("Export Error", "Please solve the problem first.")
            return
        # Allows to select a location and filename to save the Docx
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")],
            title="Save solution as..."
        )
        # If user cancels the export, aborts it
        if not file_path:
            return


        path = export_to_docx(steps_result, fig, file_path)#Exports everything into the path
        messagebox.showinfo("Export Success", f"Solution exported to:\n{path}") #Export completion pop up

    #Buttons
    solve_btn = tk.Button(parent_frame, text="Solve", font=("Arial", 18), command=on_solve)
    solve_btn.grid(row=5, column=0, pady=10, columnspan=2, sticky="ew", padx=5)

    export_btn = tk.Button(parent_frame, text="Save to Word", font=("Arial", 18), command=on_export)
    export_btn.grid(row=5, column=2, pady=10, sticky="ew", padx=5)

    back_btn = tk.Button(parent_frame, text="Back to Main", font=("Arial", 18), command=lambda: go_back(parent_frame))
    back_btn.grid(row=5, column=3, pady=10, sticky="ew", padx=5)
